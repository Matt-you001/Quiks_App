from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT_DIR = ROOT / "documents"
OUT_DIR.mkdir(exist_ok=True)
OUT_PATH = OUT_DIR / "Quiks_School_Partnership_Proposal.docx"

NAVY = "102A43"
TEAL = "087F8C"
PURPLE = "7B2CBF"
PINK = "EC5AA5"
GOLD = "F2B84B"
INK = "17212B"
MUTED = "536779"
PALE_TEAL = "E9F7F6"
PALE_PURPLE = "F4ECFA"
PALE_GOLD = "FFF7E5"
WHITE = "FFFFFF"
LINE = "D6E2E8"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=120, start=150, bottom=120, end=150):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths_dxa, indent=150):
    total = sum(widths_dxa)
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(total))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent))
    tbl_ind.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for index, cell in enumerate(row.cells):
            width = widths_dxa[index]
            cell.width = Inches(width / 1440)
            tc_w = cell._tc.get_or_add_tcPr().find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                cell._tc.get_or_add_tcPr().append(tc_w)
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_cell_border(cell, color=LINE, size=6):
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "start", "bottom", "end", "insideH", "insideV"):
        tag = f"w:{edge}"
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), str(size))
        element.set(qn("w:color"), color)


def set_font(run, size=11, color=INK, bold=False, italic=False, name="Aptos"):
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    run.bold = bold
    run.italic = italic


def add_text(doc, text, *, size=11, color=INK, bold=False, italic=False, align=None,
             before=0, after=8, line=1.15, keep=False):
    paragraph = doc.add_paragraph()
    if align is not None:
        paragraph.alignment = align
    fmt = paragraph.paragraph_format
    fmt.space_before = Pt(before)
    fmt.space_after = Pt(after)
    fmt.line_spacing = line
    fmt.keep_with_next = keep
    set_font(paragraph.add_run(text), size=size, color=color, bold=bold, italic=italic)
    return paragraph


def add_heading(doc, text, level=1):
    paragraph = doc.add_paragraph(style=f"Heading {level}")
    paragraph.paragraph_format.keep_with_next = True
    paragraph.add_run(text)
    return paragraph


def add_bullet(doc, lead, detail, *, compact=False):
    paragraph = doc.add_paragraph(style="List Bullet")
    paragraph.paragraph_format.space_after = Pt(3 if compact else 5)
    paragraph.paragraph_format.line_spacing = 1.15
    first = paragraph.add_run(lead)
    set_font(first, size=10.5, color=NAVY, bold=True)
    second = paragraph.add_run(detail)
    set_font(second, size=10.5, color=INK)
    return paragraph


def add_number(doc, number, title, detail):
    paragraph = doc.add_paragraph(style="List Number")
    paragraph.paragraph_format.space_after = Pt(5)
    paragraph.paragraph_format.line_spacing = 1.15
    title_run = paragraph.add_run(title)
    set_font(title_run, size=10.5, color=NAVY, bold=True)
    detail_run = paragraph.add_run(detail)
    set_font(detail_run, size=10.5, color=INK)
    return paragraph


def add_callout(doc, label, message, fill=PALE_TEAL, accent=TEAL):
    table = doc.add_table(rows=1, cols=1)
    set_table_geometry(table, [9360])
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    set_cell_border(cell, color=accent, size=10)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.15
    set_font(p.add_run(f"{label}  "), size=10.5, color=accent, bold=True)
    set_font(p.add_run(message), size=10.5, color=INK)
    doc.add_paragraph().paragraph_format.space_after = Pt(1)


def add_page_break(doc):
    doc.add_page_break()


def add_page_label(doc, text):
    add_text(doc, text.upper(), size=8.5, color=TEAL, bold=True, after=5, keep=True)


doc = Document()
section = doc.sections[0]
section.page_width = Inches(8.5)
section.page_height = Inches(11)
section.top_margin = Inches(0.72)
section.bottom_margin = Inches(0.72)
section.left_margin = Inches(1.0)
section.right_margin = Inches(1.0)
section.header_distance = Inches(0.35)
section.footer_distance = Inches(0.35)

styles = doc.styles
normal = styles["Normal"]
normal.font.name = "Aptos"
normal._element.rPr.rFonts.set(qn("w:ascii"), "Aptos")
normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Aptos")
normal.font.size = Pt(10.5)
normal.font.color.rgb = RGBColor.from_string(INK)
normal.paragraph_format.space_after = Pt(7)
normal.paragraph_format.line_spacing = 1.2

for level, size, before, after in ((1, 17, 12, 6), (2, 13, 8, 4), (3, 11, 6, 3)):
    style = styles[f"Heading {level}"]
    style.font.name = "Aptos Display"
    style._element.rPr.rFonts.set(qn("w:ascii"), "Aptos Display")
    style._element.rPr.rFonts.set(qn("w:hAnsi"), "Aptos Display")
    style.font.size = Pt(size)
    style.font.bold = True
    style.font.color.rgb = RGBColor.from_string(NAVY if level == 1 else TEAL)
    style.paragraph_format.space_before = Pt(before)
    style.paragraph_format.space_after = Pt(after)
    style.paragraph_format.keep_with_next = True

for style_name in ("List Bullet", "List Number"):
    style = styles[style_name]
    style.font.name = "Aptos"
    style.font.size = Pt(10.5)
    style.paragraph_format.left_indent = Inches(0.375)
    style.paragraph_format.first_line_indent = Inches(-0.194)
    style.paragraph_format.space_after = Pt(4)
    style.paragraph_format.line_spacing = 1.208

# Running header and footer.
header = section.header
header_p = header.paragraphs[0]
header_p.alignment = WD_ALIGN_PARAGRAPH.LEFT
header_p.paragraph_format.space_after = Pt(0)
set_font(header_p.add_run("QUIKS  |  SCHOOL PARTNERSHIP PROPOSAL"), size=8.5, color=MUTED, bold=True)

footer = section.footer
footer_p = footer.paragraphs[0]
footer_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
set_font(footer_p.add_run("Tech Solution Providers Ltd  |  "), size=8, color=MUTED)
field = OxmlElement("w:fldSimple")
field.set(qn("w:instr"), "PAGE")
footer_p._p.append(field)

# PAGE 1 — opening case.
logo_table = doc.add_table(rows=1, cols=3)
set_table_geometry(logo_table, [3120, 3120, 3120])
logos = [
    ("quiks-children-playstore-icon-512.png", "QUIKS CHILDREN", PURPLE),
    ("quiks-teens-playstore-icon-512.png", "QUIKS TEENS", TEAL),
    ("quiks-uni-playstore-icon-512.png", "QUIKS UNI", NAVY),
]
for cell, (filename, label, color) in zip(logo_table.rows[0].cells, logos):
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(2)
    picture = p.add_run().add_picture(str(ROOT / "assets" / "images" / filename), width=Inches(0.58))
    picture._inline.docPr.set("title", f"{label.title()} logo")
    picture._inline.docPr.set("descr", f"App logo for {label.title()}")
    p2 = cell.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p2.paragraph_format.space_after = Pt(0)
    set_font(p2.add_run(label), size=7.5, color=color, bold=True)

add_text(doc, "A Smarter Learning Partnership for Your Students", size=25, color=NAVY, bold=True,
         align=WD_ALIGN_PARAGRAPH.CENTER, before=12, after=5, line=1.0)
add_text(doc, "Proposal to [SCHOOL NAME]", size=14, color=TEAL, bold=True,
         align=WD_ALIGN_PARAGRAPH.CENTER, after=3)
add_text(doc, "Prepared by Tech Solution Providers Ltd  |  [DATE]", size=9.5, color=MUTED,
         align=WD_ALIGN_PARAGRAPH.CENTER, after=14)

add_callout(
    doc,
    "PARTNERSHIP PURPOSE",
    "To give students structured, engaging practice beyond the classroom while giving teachers practical tools to organise learning, assign work and monitor participation.",
)

add_heading(doc, "Why Quiks", 1)
add_text(
    doc,
    "Quiks is a learning and practice platform built for different stages of education: Quiks Children for younger learners, Quiks Teens for secondary-school and college-age students, and Quiks Uni for tertiary learners. For a school partnership, the appropriate variant can be recommended by age group so that content, language and challenge level remain relevant.",
    after=6,
)
add_text(
    doc,
    "The platform complements teaching rather than replacing it. Students can practise subjects, revise difficult topics, receive explanations and extend learning through the Learning Hub. Teachers can use Classroom tools to organise students, create learning activities and review submitted work. Quiks is available on mobile and web, with signed-in profiles and subscription status designed to remain consistent across devices within each variant.",
    after=7,
)

add_heading(doc, "The value to the school", 2)
add_bullet(doc, "More purposeful independent study. ", "Students have a guided route for practice, revision and topic exploration outside lesson time.")
add_bullet(doc, "Better continuity between school and home. ", "Classroom activities and learner profiles help carry assigned learning beyond the physical classroom.")
add_bullet(doc, "Motivation without removing academic focus. ", "Progression, certificates, competitions and controlled Breather activities reward effort while keeping learning central.")
add_bullet(doc, "A scalable support layer. ", "One platform can serve different ages, subjects, grades, curricula and target examinations.")

# PAGE 2 — feature-to-benefit map.
add_page_break(doc)
add_page_label(doc, "2  |  Features and educational benefits")
add_heading(doc, "What students and teachers gain", 1)
add_text(doc, "Quiks combines guided practice, instructional support and school-facing tools in one learner journey.", after=7)

feature_rows = [
    ("Adaptive practice", "Subject, grade, topic, difficulty, curriculum and target-exam context guide question generation.", "More relevant practice and clearer progression.", "A flexible resource for reinforcement, revision and differentiated work."),
    ("Training and revision", "Correct-answer explanations, Next Question and Learn More links connect practice to deeper study.", "Students understand why an answer is correct, not only whether it is correct.", "Misconceptions can be addressed through explanation-led follow-up."),
    ("Learning Hub", "Learners request lessons by subject, topic and grade, copy content, or ask a direct question.", "Supports curiosity, homework research and self-directed learning.", "Provides an additional explanation channel for students who need more support."),
    ("Classroom", "Teachers create classes, invite students by code or link, set activities and review submissions.", "Students receive structured tasks within a familiar learning app.", "Brings organisation, assignment delivery and participation into one workflow."),
    ("Competition", "One-to-one and Group Competition modes use timed challenges, codes and invitation links.", "Makes retrieval practice social, engaging and goal-oriented.", "Useful for revision events, subject clubs and friendly inter-class challenges."),
    ("Progress and certificates", "Twenty levels per grade support staged progression; grade completion can earn a branded certificate reflecting excellence and speed.", "Visible milestones reward persistence and achievement.", "Teachers and families gain a clearer conversation starter around effort and progress."),
    ("Breather activities", "After every three successful sessions, learners may take one short text or lightweight game-based Breather before returning to learning.", "Helps relieve fatigue while preserving study discipline.", "Encourages sustainable study habits without creating an unrestricted escape route."),
]

table = doc.add_table(rows=1, cols=4)
set_table_geometry(table, [1550, 2950, 2430, 2430])
headers = ["FEATURE", "WHAT IT DOES", "STUDENT BENEFIT", "TEACHER / SCHOOL BENEFIT"]
for cell, label in zip(table.rows[0].cells, headers):
    set_cell_shading(cell, NAVY)
    set_cell_border(cell, color=WHITE, size=5)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    set_font(p.add_run(label), size=8.2, color=WHITE, bold=True)
set_repeat_table_header(table.rows[0])

for index, row_data in enumerate(feature_rows):
    row = table.add_row()
    fill = "FFFFFF" if index % 2 == 0 else "F5F8FA"
    for col, (cell, value) in enumerate(zip(row.cells, row_data)):
        set_cell_shading(cell, fill)
        set_cell_border(cell)
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing = 1.05
        set_font(p.add_run(value), size=8.3, color=NAVY if col == 0 else INK, bold=(col == 0))

add_text(doc, "Quiks also includes multiple learner profiles for subscribed households, a local question library for continued practice, and free-to-paid upgrade controls for premium services.", size=9.2, color=MUTED, italic=True, before=7, after=0)

# PAGE 3 — classroom workflow and rollout.
add_page_break(doc)
add_page_label(doc, "3  |  How a school can use Quiks")
add_heading(doc, "A practical companion to teaching", 1)
add_text(doc, "Schools can introduce Quiks as an approved supplementary learning tool for homework, revision, intervention and enrichment. Teachers retain responsibility for curriculum delivery and can choose where Quiks adds the most value.", after=7)

workflow = [
    ("Define the learning focus. ", "The teacher selects a subject, topic, grade or examination objective that matches current classroom work."),
    ("Create and share the activity. ", "Using Classroom, the teacher organises the class and shares access by code or invitation link."),
    ("Students practise and learn. ", "Learners complete activities, use explanations, and open Learn More when a concept needs deeper attention."),
    ("Review and respond. ", "The teacher reviews participation and submitted work, then uses the evidence to reinforce, reteach or extend learning."),
]
for idx, (title, detail) in enumerate(workflow, start=1):
    add_number(doc, idx, title, detail)

add_heading(doc, "Suggested uses across the school", 2)
uses = doc.add_table(rows=1, cols=3)
set_table_geometry(uses, [2200, 3580, 3580])
for cell, label in zip(uses.rows[0].cells, ("USE CASE", "HOW QUIKS SUPPORTS IT", "POSSIBLE SCHOOL ROUTINE")):
    set_cell_shading(cell, TEAL)
    set_cell_border(cell, color=WHITE, size=5)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    set_font(p.add_run(label), size=8.5, color=WHITE, bold=True)
set_repeat_table_header(uses.rows[0])
use_rows = [
    ("Homework & consolidation", "Topic-focused practice plus explanations and Learn More.", "Set one focused activity after a taught unit."),
    ("Exam preparation", "Target-exam context, revision sessions and timed competition.", "Run weekly practice around priority subjects."),
    ("Intervention", "Grade-appropriate practice and repeatable training sessions.", "Assign additional guided work to learners needing support."),
    ("Enrichment", "Learning Hub lessons, direct questions and Group Competition.", "Use in clubs, holiday learning or high-achiever programmes."),
]
for index, row_data in enumerate(use_rows):
    row = uses.add_row()
    for col, (cell, value) in enumerate(zip(row.cells, row_data)):
        set_cell_shading(cell, PALE_TEAL if index % 2 else WHITE)
        set_cell_border(cell)
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing = 1.05
        set_font(p.add_run(value), size=8.8, color=NAVY if col == 0 else INK, bold=(col == 0))

add_heading(doc, "Responsible adoption", 2)
add_bullet(doc, "Teacher-guided use. ", "Quiks should complement the school's scheme of work, approved resources and professional judgement.", compact=True)
add_bullet(doc, "Account-based continuity. ", "Students should use the same signed-in account on their devices so profiles, records and subscription access can sync within the chosen variant.", compact=True)
add_bullet(doc, "Age-appropriate supervision. ", "The school and parents should agree expectations for devices, screen time, acceptable use and AI-assisted learning.", compact=True)
add_bullet(doc, "Review before scale. ", "A short pilot allows the school to validate content fit, workflows and safeguarding requirements before wider recommendation.", compact=True)

# PAGE 4 — pilot and ask.
add_page_break(doc)
add_page_label(doc, "4  |  Partnership recommendation")
add_heading(doc, "Start with a focused school pilot", 1)
add_text(doc, "We propose a time-bound pilot with a defined learner group and a small number of teacher champions. The aim is to establish usability, curriculum fit and practical value before a school-wide recommendation.", after=8)

pilot = doc.add_table(rows=1, cols=2)
set_table_geometry(pilot, [2750, 6610])
pilot_rows = [
    ("Recommended scope", "One year group or selected class; two to four subjects; nominated teacher coordinators."),
    ("Suggested duration", "Four to eight weeks, agreed with the school calendar."),
    ("Onboarding", "Account setup guidance, learner profile creation, Classroom setup and a short teacher orientation."),
    ("Success measures", "Activation, participation, completion, teacher usability feedback and learner engagement feedback."),
    ("Commercial terms", "School recommendation, sponsored access or institutional arrangements to be agreed after scoping."),
]
for index, (label, value) in enumerate(pilot_rows):
    row = pilot.rows[0] if index == 0 else pilot.add_row()
    label_cell, value_cell = row.cells
    set_cell_shading(label_cell, NAVY)
    set_cell_shading(value_cell, "FFFFFF" if index % 2 == 0 else "F5F8FA")
    set_cell_border(label_cell, color=WHITE, size=5)
    set_cell_border(value_cell)
    p1 = label_cell.paragraphs[0]
    p1.paragraph_format.space_after = Pt(0)
    set_font(p1.add_run(label.upper()), size=8.6, color=WHITE, bold=True)
    p2 = value_cell.paragraphs[0]
    p2.paragraph_format.space_after = Pt(0)
    p2.paragraph_format.line_spacing = 1.08
    set_font(p2.add_run(value), size=9.2, color=INK)

add_heading(doc, "What we are asking from [SCHOOL NAME]", 2)
add_bullet(doc, "Nominate a decision-maker and teacher champion. ", "They will help define the pilot group, subjects and expected outcomes.")
add_bullet(doc, "Permit a controlled introduction to selected learners and parents. ", "Clear guidance will explain that Quiks is supplementary to classroom teaching.")
add_bullet(doc, "Provide structured feedback. ", "Teacher and learner observations will guide the decision on broader recommendation or adoption.")

add_callout(
    doc,
    "PROPOSED NEXT STEP",
    "Schedule a 30-minute demonstration and scoping conversation to review the appropriate Quiks variant, the school's priority use case and a suitable pilot group.",
    fill=PALE_PURPLE,
    accent=PURPLE,
)

add_heading(doc, "Contact", 2)
contact = doc.add_table(rows=3, cols=2)
set_table_geometry(contact, [2200, 7160])
contact_data = [
    ("Organisation", "Tech Solution Providers Ltd"),
    ("Representative", "[NAME AND TITLE]"),
    ("Email / telephone", "[EMAIL ADDRESS]  |  [TELEPHONE NUMBER]"),
]
for row, (label, value) in zip(contact.rows, contact_data):
    set_cell_shading(row.cells[0], PALE_GOLD)
    set_cell_shading(row.cells[1], WHITE)
    for cell in row.cells:
        set_cell_border(cell)
    p1 = row.cells[0].paragraphs[0]
    p1.paragraph_format.space_after = Pt(0)
    set_font(p1.add_run(label), size=9, color=NAVY, bold=True)
    p2 = row.cells[1].paragraphs[0]
    p2.paragraph_format.space_after = Pt(0)
    set_font(p2.add_run(value), size=9.2, color=INK)

add_text(doc, "Quiks: Learn fast. Grow steadily.", size=11, color=TEAL, bold=True,
         align=WD_ALIGN_PARAGRAPH.CENTER, before=10, after=0)

# Metadata.
props = doc.core_properties
props.title = "Quiks School Partnership Proposal"
props.subject = "Proposal for schools to recommend Quiks to students"
props.author = "Tech Solution Providers Ltd"
props.keywords = "Quiks, school partnership, educational technology, student learning, teachers"

doc.save(OUT_PATH)
print(OUT_PATH)
